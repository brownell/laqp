"""
Louisiana QSO Party - Individual Results Generator

Generates individual DOCX files for each contestant with:
- Overall placement (among all logs)
- Category placement (within their category)
- Score breakdown
- Parishes/States/Provinces/DX worked
- QSOs by band and mode

VERSION 2: Fixed spacing and label formatting
"""
import sys
from pathlib import Path
from typing import Dict, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from config.config import (
    CONTEST_NAME, CONTEST_YEAR, SPONSOR_NAME, SPONSOR_WEBSITE,
    INDIVIDUAL_RESULTS_DIR
)
from laqp.categories import CATEGORIES


class IndividualResultsGenerator:
    """Generates individual DOCX result files for contestants"""
    
    def __init__(self, contest_year: int = None):
        """
        Initialize the generator.
        
        Args:
            contest_year: Year of the contest (defaults to config.CONTEST_YEAR)
        """
        self.contest_year = contest_year or CONTEST_YEAR
        self.contest_name = f"{CONTEST_NAME} {self.contest_year}"
    
    def generate_individual_result(self, 
                                   callsign: str,
                                   category_short: str,
                                   overlay: str,
                                   overall_placement: int,
                                   total_logs: int,
                                   category_placement: int,
                                   category_total: int,
                                   final_score: int,
                                   total_qsos: int,
                                   multipliers: int,
                                   parishes_worked: int,
                                   states_worked: int,
                                   provinces_worked: int,
                                   dx_contacts: int,
                                   qsos_by_band: Dict[int, int],
                                   qsos_by_mode: Dict[str, int],
                                   bands_worked: List[int],
                                   output_dir: Path = None) -> Path:
        """
        Generate an individual result DOCX file for a contestant.
        
        Args:
            callsign: Station callsign
            category_short: Short category name (e.g., 'nl_ph_lo')
            overlay: Overlay type ('WIRES', 'TB-WIRES', 'POTA', or None)
            overall_placement: Placement among all logs (1-based)
            total_logs: Total number of logs in contest
            category_placement: Placement within category (1-based)
            category_total: Total logs in category
            final_score: Final calculated score
            total_qsos: Total number of QSOs
            multipliers: Total multipliers
            parishes_worked: Number of parishes worked
            states_worked: Number of states worked
            provinces_worked: Number of provinces worked
            dx_contacts: Number of DX contacts
            qsos_by_band: Dict of {band: count}
            qsos_by_mode: Dict of {mode: count}
            bands_worked: List of bands worked
            output_dir: Output directory (defaults to INDIVIDUAL_RESULTS_DIR)
        
        Returns:
            Path to generated DOCX file
        """
        if output_dir is None:
            output_dir = INDIVIDUAL_RESULTS_DIR
        
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Create document
        doc = Document()
        
        # Set default font and reduce spacing
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Reduce paragraph spacing globally
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.0
        
        # ===== HEADER =====
        self._add_header(doc, callsign)
        
        # ===== CATEGORY INFO =====
        self._add_category_info(doc, category_short, overlay)
        
        # ===== FINAL RESULTS =====
        self._add_results_section(
            doc, overall_placement, total_logs, category_placement,
            category_total, final_score, total_qsos, multipliers
        )
        
        # ===== BREAKDOWN =====
        self._add_breakdown_section(
            doc, parishes_worked, states_worked, provinces_worked, dx_contacts
        )
        
        # ===== QSOs BY BAND =====
        self._add_band_breakdown(doc, qsos_by_band, bands_worked)
        
        # ===== QSOs BY MODE =====
        self._add_mode_breakdown(doc, qsos_by_mode)
        
        # ===== FOOTER =====
        self._add_footer(doc)
        
        # Save document
        filename = f"{callsign.upper()}.docx"
        filepath = output_dir / filename
        doc.save(str(filepath))
        
        return filepath
    
    def _reduce_spacing(self, paragraph):
        """Reduce spacing for a paragraph"""
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.0
    
    def _add_header(self, doc: Document, callsign: str):
        """Add document header"""
        # Title
        heading = doc.add_heading(self.contest_name.upper(), level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._reduce_spacing(heading)
        
        # Subtitle
        subtitle = doc.add_paragraph('INDIVIDUAL RESULTS')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].font.bold = True
        self._reduce_spacing(subtitle)
        
        # Separator
        sep = doc.add_paragraph('=' * 60)
        self._reduce_spacing(sep)
        
        # Callsign (label bold, value not bold)
        callsign_para = doc.add_paragraph()
        callsign_para.add_run('Callsign: ').font.bold = True
        callsign_para.add_run(callsign.upper())
        self._reduce_spacing(callsign_para)
    
    def _add_category_info(self, doc: Document, category_short: str, overlay: str):
        """Add category information"""
        category_full = CATEGORIES.get(category_short, 'Unknown Category')
        
        # Category (label bold, value not bold)
        para = doc.add_paragraph()
        para.add_run('Category: ').font.bold = True
        para.add_run(category_full)
        self._reduce_spacing(para)
        
        # Overlay info (label bold, value not bold)
        if overlay:
            para = doc.add_paragraph()
            para.add_run('Overlay: ').font.bold = True
            para.add_run(overlay)
            self._reduce_spacing(para)
        else:
            para = doc.add_paragraph('(No overlay)')
            para.runs[0].font.italic = True
            self._reduce_spacing(para)
    
    def _add_results_section(self, doc: Document, overall_placement: int,
                            total_logs: int, category_placement: int,
                            category_total: int, final_score: int,
                            total_qsos: int, multipliers: int):
        """Add final results section"""
        # Section heading
        heading = doc.add_heading('FINAL RESULTS', level=2)
        self._reduce_spacing(heading)
        
        # Overall placement (label bold, value not bold)
        para = doc.add_paragraph()
        para.add_run('Overall Placement: ').font.bold = True
        rank_text = self._ordinal(overall_placement)
        para.add_run(f"{rank_text} overall (out of {total_logs} logs)")
        self._reduce_spacing(para)
        
        # Category placement (label bold, value not bold)
        para = doc.add_paragraph()
        para.add_run('Category Placement: ').font.bold = True
        rank_text = self._ordinal(category_placement)
        para.add_run(f"{rank_text} in category (out of {category_total} logs)")
        self._reduce_spacing(para)
        
        # Score (label bold, value not bold)
        para = doc.add_paragraph()
        para.add_run('Final Score: ').font.bold = True
        para.add_run(f"{final_score:,} points")
        self._reduce_spacing(para)
        
        # QSOs (label bold, value not bold)
        para = doc.add_paragraph()
        para.add_run('Total QSOs: ').font.bold = True
        para.add_run(str(total_qsos))
        self._reduce_spacing(para)
        
        # Multipliers (label bold, value not bold)
        para = doc.add_paragraph()
        para.add_run('Multipliers: ').font.bold = True
        para.add_run(str(multipliers))
        self._reduce_spacing(para)
    
    def _add_breakdown_section(self, doc: Document, parishes: int,
                               states: int, provinces: int, dx: int):
        """Add breakdown section"""
        heading = doc.add_heading('BREAKDOWN', level=2)
        self._reduce_spacing(heading)
        
        # All labels bold, values not bold
        para = doc.add_paragraph()
        para.add_run('Parishes Worked: ').font.bold = True
        para.add_run(str(parishes))
        self._reduce_spacing(para)
        
        para = doc.add_paragraph()
        para.add_run('States Worked: ').font.bold = True
        para.add_run(str(states))
        self._reduce_spacing(para)
        
        para = doc.add_paragraph()
        para.add_run('Provinces Worked: ').font.bold = True
        para.add_run(str(provinces))
        self._reduce_spacing(para)
        
        para = doc.add_paragraph()
        para.add_run('DX Contacts: ').font.bold = True
        para.add_run(str(dx))
        self._reduce_spacing(para)
    
    def _add_band_breakdown(self, doc: Document, qsos_by_band: Dict[int, int],
                           bands_worked: List[int]):
        """Add QSOs by band section"""
        heading = doc.add_heading('QSOs BY BAND', level=2)
        self._reduce_spacing(heading)
        
        # Sort bands by frequency (160m first, then 80, 40, etc.)
        band_order = [160, 80, 40, 20, 15, 10, 6, 2]
        
        for band in band_order:
            if band in bands_worked:
                count = qsos_by_band.get(band, 0)
                # Band label bold, count not bold
                para = doc.add_paragraph()
                para.add_run(f'  {band}m: ').font.bold = True
                para.add_run(str(count))
                self._reduce_spacing(para)
    
    def _add_mode_breakdown(self, doc: Document, qsos_by_mode: Dict[str, int]):
        """Add QSOs by mode section"""
        heading = doc.add_heading('QSOs BY MODE', level=2)
        self._reduce_spacing(heading)
        
        # Group modes
        phone_count = qsos_by_mode.get('Phone', 0)
        cw_digital_count = qsos_by_mode.get('CW/Digital', 0)
        
        # Mode label bold, count not bold
        para = doc.add_paragraph()
        para.add_run('  Phone: ').font.bold = True
        para.add_run(str(phone_count))
        self._reduce_spacing(para)
        
        para = doc.add_paragraph()
        para.add_run('  CW/Digital: ').font.bold = True
        para.add_run(str(cw_digital_count))
        self._reduce_spacing(para)
    
    def _add_footer(self, doc: Document):
        """Add footer"""
        sep = doc.add_paragraph('=' * 60)
        self._reduce_spacing(sep)
        
        footer = doc.add_paragraph(
            f'Thank you for participating in the {self.contest_name}!'
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._reduce_spacing(footer)
        
        sponsor = doc.add_paragraph(
            f'Sponsored by {SPONSOR_NAME} - {SPONSOR_WEBSITE}'
        )
        sponsor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sponsor.runs[0].font.size = Pt(10)
        self._reduce_spacing(sponsor)
    
    def _ordinal(self, n: int) -> str:
        """Convert number to ordinal (1st, 2nd, 3rd, etc.)"""
        if 10 <= n % 100 <= 20:
            suffix = 'th'
        else:
            suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th')
        return f"{n}{suffix}"


def generate_all_individual_results(score_data: List[Dict],
                                    category_standings: Dict[str, List[Dict]],
                                    output_dir: Path = None) -> List[Path]:
    """
    Generate individual result files for all contestants.

    Adapted from TQP statistics.py for LA rules created by Charles Sanders, NO5W
    
    Args:
        score_data: List of dicts with score info for each log
        category_standings: Dict of {category: [logs sorted by score]}
        output_dir: Output directory
    
    Returns:
        List of paths to generated files
    """
    generator = IndividualResultsGenerator()
    generated_files = []
    
    # Sort score_data by score for overall standings
    sorted_scores = sorted(score_data, key=lambda x: x['final_score'], reverse=True)
    total_logs = len(sorted_scores)
    
    # Process each log
    for overall_idx, log_data in enumerate(sorted_scores, 1):
        callsign = log_data['callsign']
        category = log_data['category']
        
        # Find category placement
        category_logs = category_standings.get(category, [])
        category_total = len(category_logs)
        
        # Find this log's position in category
        category_placement = 1
        for idx, cat_log in enumerate(category_logs, 1):
            if cat_log['callsign'] == callsign:
                category_placement = idx
                break
        
        # Generate individual file
        filepath = generator.generate_individual_result(
            callsign=callsign,
            category_short=category,
            overlay=log_data.get('overlay'),
            overall_placement=overall_idx,
            total_logs=total_logs,
            category_placement=category_placement,
            category_total=category_total,
            final_score=log_data['final_score'],
            total_qsos=log_data['total_qsos'],
            multipliers=log_data['multipliers'],
            parishes_worked=log_data['parishes_worked'],
            states_worked=log_data['states_worked'],
            provinces_worked=log_data['provinces_worked'],
            dx_contacts=log_data['dx_contacts'],
            qsos_by_band=log_data['qsos_by_band'],
            qsos_by_mode=log_data['qsos_by_mode'],
            bands_worked=log_data['bands_worked'],
            output_dir=output_dir
        )
        
        generated_files.append(filepath)
        print(f"  Generated: {filepath.name}")
    
    return generated_files


if __name__ == "__main__":
    print("Individual Results Generator")
    print("This module should be imported, not run directly.")
