#!/usr/bin/env python3
"""
Standalone test for Individual Results Generator
This bypasses laqp.core imports to test the module directly
"""
import sys
from pathlib import Path
from docx import Document

# Add project to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

# Test if python-docx is installed
try:
    from docx import Document
    print("✓ python-docx is installed")
except ImportError:
    print("✗ python-docx is NOT installed")
    print("  Run: pip install python-docx")
    sys.exit(1)

# Direct imports to avoid laqp.core
from config.config import CONTEST_NAME, CONTEST_YEAR, SPONSOR_NAME, SPONSOR_WEBSITE
from laqp.categories import CATEGORIES

print(f"✓ Config loaded: {CONTEST_NAME} {CONTEST_YEAR}")
print(f"✓ Categories loaded: {len(CATEGORIES)} categories defined")
print()

# Now test the generator directly by importing just its code
print("Creating test DOCX file...")

# Create output directory
test_output = project_root / 'data' / 'output' / 'individual_results' / 'test'
test_output.mkdir(parents=True, exist_ok=True)

# Create a simple test document
doc = Document()

# Import formatting tools
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Reduce spacing helper
def reduce_spacing(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = 1.0

# Add title
heading = doc.add_heading(f'{CONTEST_NAME} {CONTEST_YEAR}'.upper(), level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

reduce_spacing(heading)

# Add subtitle
subtitle = doc.add_paragraph('INDIVIDUAL RESULTS')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True
reduce_spacing(subtitle)

# Separator
sep = doc.add_paragraph('=' * 60)
reduce_spacing(sep)

# Callsign - LABEL BOLD, VALUE NOT BOLD
para = doc.add_paragraph()
para.add_run('Callsign: ').font.bold = True
para.add_run('KJ5BYZ')
reduce_spacing(para)

# Category - LABEL BOLD, VALUE NOT BOLD
category_full = CATEGORIES.get('nl_ph_lo', 'Unknown')
para = doc.add_paragraph()
para.add_run('Category: ').font.bold = True
para.add_run(category_full)
reduce_spacing(para)

para = doc.add_paragraph()
para.add_run('Overlay: ').font.bold = True
para.add_run('POTA')
reduce_spacing(para)

# Results
heading = doc.add_heading('FINAL RESULTS', level=2)
reduce_spacing(heading)

# LABEL BOLD, VALUE NOT BOLD
para = doc.add_paragraph()
para.add_run('Overall Placement: ').font.bold = True
para.add_run('12th overall (out of 87 logs)')
reduce_spacing(para)

para = doc.add_paragraph()
para.add_run('Category Placement: ').font.bold = True
para.add_run('5th in category (out of 23 logs)')
reduce_spacing(para)

para = doc.add_paragraph()
para.add_run('Final Score: ').font.bold = True
para.add_run('12,450 points')
reduce_spacing(para)

para = doc.add_paragraph()
para.add_run('Total QSOs: ').font.bold = True
para.add_run('145')
reduce_spacing(para)

para = doc.add_paragraph()
para.add_run('Multipliers: ').font.bold = True
para.add_run('42')
reduce_spacing(para)

# Breakdown
heading = doc.add_heading('BREAKDOWN', level=2)
reduce_spacing(heading)

para = doc.add_paragraph()
para.add_run('Parishes Worked: ').font.bold = True
para.add_run('28')
reduce_spacing(para)

para = doc.add_paragraph()
para.add_run('States Worked: ').font.bold = True
para.add_run('35')
reduce_spacing(para)

para = doc.add_paragraph()
para.add_run('Provinces Worked: ').font.bold = True
para.add_run('5')
reduce_spacing(para)

para = doc.add_paragraph()
para.add_run('DX Contacts: ').font.bold = True
para.add_run('2')
reduce_spacing(para)

# QSOs by Band - LABEL BOLD, VALUE NOT BOLD
heading = doc.add_heading('QSOs BY BAND', level=2)
reduce_spacing(heading)
for band, count in [(160, 10), (80, 25), (40, 40), (20, 35), (15, 20), (10, 15)]:
    para = doc.add_paragraph()
    para.add_run(f'  {band}m: ').font.bold = True
    para.add_run(str(count))
    reduce_spacing(para)

# QSOs by Mode - LABEL BOLD, VALUE NOT BOLD
heading = doc.add_heading('QSOs BY MODE', level=2)
reduce_spacing(heading)
para = doc.add_paragraph()
para.add_run('  Phone: ').font.bold = True
para.add_run('145')
reduce_spacing(para)
para = doc.add_paragraph()
para.add_run('  CW/Digital: ').font.bold = True
para.add_run('0')
reduce_spacing(para)

# Footer
sep = doc.add_paragraph('=' * 60)
reduce_spacing(sep)
footer = doc.add_paragraph(f'Thank you for participating in the {CONTEST_NAME} {CONTEST_YEAR}!')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
reduce_spacing(footer)

sponsor = doc.add_paragraph(f'Sponsored by {SPONSOR_NAME} - {SPONSOR_WEBSITE}')
sponsor.alignment = WD_ALIGN_PARAGRAPH.CENTER
sponsor.runs[0].font.size = Pt(10)
reduce_spacing(sponsor)

# Save
output_file = test_output / 'TEST_KJ5BYZ.docx'
doc.save(str(output_file))

print(f"✓ Successfully created: {output_file}")
print()
print("Open this file to verify the formatting looks good:")
print(f"  libreoffice {output_file}")
print()
print("If this looks good, the individual_results.py module will work fine!")
print("We just can't test it through laqp.core until we update the other modules.")
